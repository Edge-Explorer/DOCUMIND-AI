# ingest/api.py

import os
import re
import tempfile
import subprocess
import shutil
import logging
from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
import fitz  # PyMuPDF for PDF fallback
from pdfminer.high_level import extract_text as extract_text_pdfminer
import docx

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask Blueprint
ingest_bp = Blueprint('ingest', __name__)

# Directories for storing documents and index
DOCUMENTS_FOLDER = os.path.join(os.getcwd(), 'documents')
INDEX_DIR = os.path.join(os.getcwd(), 'indices')
OCR_FOLDER = os.path.join(os.getcwd(), 'ocr_documents')  # New folder for OCR'd versions

# Ensure directories exist
os.makedirs(DOCUMENTS_FOLDER, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)
os.makedirs(OCR_FOLDER, exist_ok=True)  # Ensure OCR folder exists

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc'}

def allowed_file(filename: str) -> bool:
    """Return True if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def is_ocr_needed(pdf_path: str) -> bool:
    """
    Check if a PDF likely needs OCR by attempting to extract text
    and evaluating the result.
    """
    try:
        # Try standard text extraction
        text = extract_text_pdfminer(pdf_path)
        
        # If we got a reasonable amount of text, OCR may not be needed
        if text and len(text.strip()) > 100:
            # Check for common OCR indicators (could be refined further)
            suspicious_patterns = [
                "ï¿½", # Unicode replacement character
                "\ufffd", # Another form of Unicode replacement
                "TCPDF", # Often appears in PDFs with poor text extraction
            ]
            
            # Check if suspicious patterns make up a significant proportion of the text
            if any(pattern in text for pattern in suspicious_patterns):
                return True
                
            # If we have text and no suspicious patterns, OCR likely not needed
            return False
        
        # Not enough text extracted, try PyMuPDF as backup
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        sample_text = ""
        
        # Get text from first 3 pages or all pages if fewer
        for i in range(min(3, page_count)):
            sample_text += doc[i].get_text() or ""
        doc.close()
        
        # If we still don't have much text, OCR is needed
        if len(sample_text.strip()) < 100:
            return True
            
        return False
    except Exception as e:
        logger.warning(f"Error checking OCR need: {e}")
        # When in doubt, perform OCR
        return True


def perform_ocr_on_pdf(pdf_path: str) -> str:
    """
    Perform OCR on a PDF file and return the path to the OCR'd version.
    """
    try:
        # Generate output filename in the OCR folder
        base_name = os.path.basename(pdf_path)
        ocr_output_path = os.path.join(OCR_FOLDER, f"{os.path.splitext(base_name)[0]}_ocr.pdf")
        
        # Use ocrmypdf to perform OCR
        cmd = ["ocrmypdf", "--force-ocr", pdf_path, ocr_output_path]
        
        logger.info(f"Running OCR command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, check=False)
        
        if result.returncode != 0:
            logger.error(f"OCR failed: {result.stderr}")
            # Fall back to original file if OCR fails
            return pdf_path
            
        logger.info(f"OCR completed successfully for {base_name}")
        return ocr_output_path
    except Exception as e:
        logger.error(f"Error during OCR process: {e}")
        # Return original path if anything goes wrong
        return pdf_path


def extract_text_from_pdf(path: str, try_ocr: bool = True) -> list:
    """
    Extract text from PDF with page tracking
    Returns a list of (page_text, page_number) tuples
    """
    result = []
    
    # Try using PyMuPDF for page-by-page extraction first
    try:
        doc = fitz.open(path)
        page_count = doc.page_count
        logger.info(f"PDF has {page_count} pages")
        
        for page_num in range(page_count):
            page = doc[page_num]
            text = page.get_text() or ''
            if text.strip():
                # Store page text with 1-based page numbering
                result.append((text, page_num + 1))
        doc.close()
        
        # If we got a reasonable amount of text, return it
        if result and sum(len(text) for text, _ in result) > 200:
            return result
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")
    
    # Fallback to pdfminer if PyMuPDF didn't work well
    if not result:
        try:
            # Unfortunately pdfminer doesn't easily support page-by-page extraction
            text = extract_text_pdfminer(path)
            if text and len(text.strip()) > 100:
                # If we can't track pages, just assign everything to page 1
                result = [(text, 1)]
        except Exception as e:
            logger.error(f"pdfminer extraction failed: {e}")
    
    # If we still don't have good results and OCR is enabled, try OCR
    if (not result or sum(len(text) for text, _ in result) < 200) and try_ocr and is_ocr_needed(path):
        logger.info(f"Standard text extraction failed or insufficient, attempting OCR for {path}")
        try:
            ocr_path = perform_ocr_on_pdf(path)
            if ocr_path != path:  # If OCR was successful and created a new file
                # Try extracting text from the OCR'd version, but don't try OCR again
                return extract_text_from_pdf(ocr_path, try_ocr=False)
        except Exception as e:
            logger.error(f"OCR process failed: {e}")
    
    # Return what we have, even if it's empty
    return result if result else [(f"Failed to extract text from {path}", 1)]


def extract_text_from_docx(path: str) -> list:
    """
    Extract text from DOCX using python-docx with page awareness
    Returns a list of (content, page_number) tuples
    """
    try:
        doc = docx.Document(path)
        
        # Word doesn't directly expose page numbers in python-docx
        # We'll estimate pages by counting paragraphs and tables
        # This is a rough approximation - around 15-20 paragraphs per page
        PARAGRAPHS_PER_PAGE = 18
        
        paragraphs = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                paragraphs.append(p.text)
        
        # Extract table data
        tables_text = []
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    table_content.append(' | '.join(cells))
            if table_content:
                tables_text.append('\n'.join(table_content))
                # Count each table as equivalent to 3 paragraphs for page estimation
                paragraphs.extend([''] * 3)
        
        # Combine all content with estimated page numbers
        result = []
        total_items = len(paragraphs)
        
        # Group paragraphs into estimated pages
        for i in range(0, total_items, PARAGRAPHS_PER_PAGE):
            page_number = (i // PARAGRAPHS_PER_PAGE) + 1
            page_content = '\n'.join(paragraphs[i:min(i+PARAGRAPHS_PER_PAGE, total_items)])
            if page_content.strip():
                result.append((page_content, page_number))
        
        return result
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return [("Failed to extract text from document", 1)]


def extract_text_from_txt(path: str) -> list:
    """
    Extract text from plain text file with estimated page breaks
    Returns a list of (content, page_number) tuples
    """
    try:
        # First try UTF-8
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Fall back to Latin-1
            with open(path, 'r', encoding='latin-1') as f:
                content = f.read()
        
        # Split into lines and estimate pages (about 40-45 lines per page)
        LINES_PER_PAGE = 40
        lines = content.split('\n')
        result = []
        
        for i in range(0, len(lines), LINES_PER_PAGE):
            page_number = (i // LINES_PER_PAGE) + 1
            page_content = '\n'.join(lines[i:i+LINES_PER_PAGE])
            if page_content.strip():
                result.append((page_content, page_number))
        
        return result
    except Exception as e:
        logger.error(f"Text file extraction failed: {e}")
        return [("Failed to extract text from document", 1)]


def create_or_update_index(docs: list[Document]) -> FAISS:
    """Load existing FAISS index or create a new one, then add documents"""
    embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
    
    index_path = os.path.join(INDEX_DIR, "index")
    
    try:
        # Try to load existing index
        db = FAISS.load_local(index_path, embeddings)
        logger.info(f"Loaded existing index from {index_path}")
        
        # Add new documents to existing index
        db.add_documents(docs)
        logger.info(f"Added {len(docs)} documents to existing index")
        
    except Exception as e:
        # If loading fails, create new index
        logger.info(f"Creating new index (load failed: {e})")
        db = FAISS.from_documents(docs, embeddings)
        logger.info(f"Created new index with {len(docs)} documents")
    
    # Save the updated index
    db.save_local(index_path)
    logger.info(f"Saved index to {index_path}")
    
    return db


@ingest_bp.route('/upload-document', methods=['POST'])
def upload_document():
    """Endpoint to upload and index a document"""
    # Validate request
    if 'document' not in request.files:
        return jsonify({'error': 'No file part in request'}), 400
    
    file = request.files['document']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': f'Invalid file type: {file.filename}'}), 400

    try:
        # Save file
        filename = secure_filename(file.filename)
        filepath = os.path.join(DOCUMENTS_FOLDER, filename)
        
        # Check if file already exists
        if os.path.exists(filepath):
            logger.warning(f"File {filename} already exists, will be overwritten")
        
        file.save(filepath)
        logger.info(f"Saved file {filename}")

        # Extract text based on extension
        ext = filename.rsplit('.', 1)[1].lower()
        
        if ext == 'pdf':
            page_texts = extract_text_from_pdf(filepath)
        elif ext in ('docx', 'doc'):
            page_texts = extract_text_from_docx(filepath)
        else:
            # txt or other plain text
            page_texts = extract_text_from_txt(filepath)

        # Basic validation
        if not page_texts or sum(len(text) for text, _ in page_texts) < 50:
            os.remove(filepath)
            return jsonify({'error': 'Failed to extract sufficient text from document'}), 400

        logger.info(f"Extracted text from {len(page_texts)} pages")

        # ✅ FIX: Split text into chunks with CORRECT metadata for each page
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        all_docs = []
        
        for text, page_number in page_texts:
            # Split this page's text into chunks
            chunks = splitter.split_text(text)
            
            # ✅ CRITICAL: Create documents with CORRECT source and page metadata
            for chunk in chunks:
                doc = Document(
                    page_content=chunk, 
                    metadata={
                        'source': filename,        # ✅ Use the actual filename
                        'page_number': page_number # ✅ Use the actual page number
                    }
                )
                all_docs.append(doc)
        
        logger.info(f"Split text into {len(all_docs)} chunks across {len(page_texts)} pages")

        # ✅ Index documents with correct metadata
        create_or_update_index(all_docs)
        
        logger.info(f"✅ Successfully indexed {filename}")
        
        return jsonify({
            'message': f'Successfully uploaded and indexed {filename}', 
            'chunks': len(all_docs),
            'pages': len(page_texts),
            'filename': filename
        }), 200
        
    except Exception as e:
        logger.error(f"Error processing upload: {e}", exc_info=True)
        
        # Clean up file if indexing failed
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
                logger.info(f"Cleaned up file {filename} after error")
            except:
                pass
        
        return jsonify({
            'error': f'Failed to process document: {str(e)}'
        }), 500


@ingest_bp.route('/delete-document', methods=['POST'])
def delete_document():
    """Endpoint to delete a document and its index entries"""
    data = request.get_json(silent=True) or {}
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400

    filepath = os.path.join(DOCUMENTS_FOLDER, filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    try:
        # Delete the file
        os.remove(filepath)
        logger.info(f"Deleted file {filename}")

        # Clean up OCR'd version if it exists
        base_name = os.path.splitext(filename)[0]
        ocr_path = os.path.join(OCR_FOLDER, f"{base_name}_ocr.pdf")
        if os.path.exists(ocr_path):
            os.remove(ocr_path)
            logger.info(f"Deleted OCR file {ocr_path}")

        # Update index - remove chunks from this document
        try:
            embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
            index_path = os.path.join(INDEX_DIR, "index")
            db = FAISS.load_local(index_path, embeddings)
            
            # Find and delete all chunks from this document
            to_delete = [
                doc_id for doc_id, doc in db.docstore._dict.items() 
                if doc.metadata.get('source') == filename
            ]
            
            if to_delete:
                db.delete(to_delete)
                db.save_local(index_path)
                logger.info(f"Deleted {len(to_delete)} chunks from index for {filename}")
            
            return jsonify({
                'message': f'Successfully deleted {filename} and removed {len(to_delete)} chunks from index'
            }), 200
            
        except Exception as e:
            logger.error(f"Index update failed: {e}")
            return jsonify({
                'warning': f'Document file deleted but failed to update index: {str(e)}'
            }), 200
            
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        return jsonify({'error': f'Failed to delete document: {str(e)}'}), 500