import os
import sys
import fitz
import docx
import logging
from typing import List, Tuple, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

DATA_DIR = "documents"
INDEX_DIR = "indices"

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}


def extract_text_from_pdf(filepath: str) -> Tuple[str, int]:
    """
    Extract text from PDF with page tracking.
    Returns: (extracted_text, page_count)
    """
    try:
        text = ""
        pdf = fitz.open(filepath)
        page_count = pdf.page_count
        
        for page_num in range(page_count):
            try:
                page = pdf[page_num]
                page_text = page.get_text()
                if page_text:
                    text += page_text
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1} from {filepath}: {e}")
                continue
        
        pdf.close()
        
        if not text.strip():
            logger.warning(f"No text extracted from PDF: {filepath}")
            return "", page_count
            
        return text, page_count
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {filepath}: {e}")
        return "", 0


def extract_text_from_docx(filepath: str) -> Tuple[str, int]:
    """
    Extract text from DOCX file.
    Returns: (extracted_text, estimated_page_count)
    """
    try:
        doc = docx.Document(filepath)
        
        # Extract paragraphs
        paragraphs = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                paragraphs.append(p.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        
        text = '\n'.join(paragraphs)
        
        if not text.strip():
            logger.warning(f"No text extracted from DOCX: {filepath}")
            return "", 1
        
        # Estimate page count (roughly 15-20 paragraphs per page)
        estimated_pages = max(1, len(paragraphs) // 18)
        
        return text, estimated_pages
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {filepath}: {e}")
        return "", 0


def extract_text_from_txt(filepath: str) -> Tuple[str, int]:
    """
    Extract text from plain text file.
    Returns: (extracted_text, estimated_page_count)
    """
    try:
        # Try UTF-8 first
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to Latin-1
            logger.info(f"UTF-8 decode failed for {filepath}, trying Latin-1")
            with open(filepath, 'r', encoding='latin-1') as f:
                text = f.read()
        
        if not text.strip():
            logger.warning(f"No text extracted from TXT: {filepath}")
            return "", 1
        
        # Estimate pages (about 40-45 lines per page)
        lines = text.split('\n')
        estimated_pages = max(1, len(lines) // 40)
        
        return text, estimated_pages
        
    except Exception as e:
        logger.error(f"Failed to extract text from TXT {filepath}: {e}")
        return "", 0


def load_documents() -> List[Tuple[str, str, int]]:
    """
    Load all documents from DATA_DIR.
    Returns: List of (text_content, filename, page_count) tuples
    """
    documents = []
    
    if not os.path.exists(DATA_DIR):
        logger.error(f"Data directory not found: {DATA_DIR}")
        return documents
    
    files = os.listdir(DATA_DIR)
    if not files:
        logger.warning(f"No files found in {DATA_DIR}")
        return documents
    
    logger.info(f"Found {len(files)} files in {DATA_DIR}")
    
    for filename in files:
        filepath = os.path.join(DATA_DIR, filename)
        
        # Skip directories
        if os.path.isdir(filepath):
            continue
        
        # Get file extension
        _, ext = os.path.splitext(filename)
        ext_lower = ext.lower()
        
        # Check if supported
        if ext_lower not in SUPPORTED_EXTENSIONS:
            logger.warning(f"Skipping unsupported file type: {filename}")
            continue
        
        logger.info(f"Processing: {filename}")
        
        # Extract text based on file type
        text = ""
        page_count = 0
        
        try:
            if ext_lower == '.pdf':
                text, page_count = extract_text_from_pdf(filepath)
            elif ext_lower in ('.docx', '.doc'):
                text, page_count = extract_text_from_docx(filepath)
            elif ext_lower == '.txt':
                text, page_count = extract_text_from_txt(filepath)
            
            # Validate extracted text
            if not text or len(text.strip()) < 50:
                logger.warning(f"Insufficient text extracted from {filename} (only {len(text)} chars)")
                continue
            
            documents.append((text, filename, page_count))
            logger.info(f"✓ Successfully loaded {filename} ({page_count} pages, {len(text)} chars)")
            
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            continue
    
    return documents


def create_document_chunks(documents: List[Tuple[str, str, int]]) -> List[Document]:
    """
    Split documents into chunks with proper metadata.
    Returns: List of Document objects with metadata
    """
    if not documents:
        logger.error("No documents to split")
        return []
    
    logger.info("Splitting documents into chunks...")
    
    # Initialize text splitter
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    all_chunks = []
    
    for text, filename, page_count in documents:
        try:
            # Split text into chunks
            text_chunks = splitter.split_text(text)
            
            # Create Document objects with metadata
            for chunk_idx, chunk_text in enumerate(text_chunks):
                # Estimate page number for this chunk
                # This is approximate - chunks are distributed across pages
                chars_per_page = len(text) / max(page_count, 1)
                chunk_position = chunk_idx * 1000  # Approximate position in document
                estimated_page = min(page_count, max(1, int(chunk_position / chars_per_page) + 1))
                
                doc = Document(
                    page_content=chunk_text,
                    metadata={
                        'source': filename,
                        'page_number': estimated_page,
                        'chunk_index': chunk_idx,
                        'total_pages': page_count
                    }
                )
                all_chunks.append(doc)
            
            logger.info(f"✓ Split {filename} into {len(text_chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Error splitting {filename}: {e}")
            continue
    
    return all_chunks


def create_index(chunks: List[Document]) -> Optional[FAISS]:
    """
    Create FAISS index from document chunks.
    Returns: FAISS index or None on failure
    """
    if not chunks:
        logger.error("No chunks to index")
        return None
    
    try:
        logger.info("Loading embedding model...")
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        logger.info(f"Creating FAISS index with {len(chunks)} chunks...")
        db = FAISS.from_documents(chunks, embeddings)
        
        return db
        
    except Exception as e:
        logger.error(f"Failed to create index: {e}")
        return None


def save_index(db: FAISS) -> bool:
    """
    Save FAISS index to disk.
    Returns: True on success, False on failure
    """
    try:
        os.makedirs(INDEX_DIR, exist_ok=True)
        index_path = os.path.join(INDEX_DIR, "index")
        
        db.save_local(index_path)
        logger.info(f"✅ Index saved successfully to: {index_path}")
        
        # Verify the index was saved
        if os.path.exists(f"{index_path}.faiss"):
            logger.info("✓ Index file verified")
            return True
        else:
            logger.error("✗ Index file not found after save")
            return False
            
    except Exception as e:
        logger.error(f"Failed to save index: {e}")
        return False


def print_summary(chunks: List[Document]):
    """Print a summary of the indexed documents"""
    if not chunks:
        return
    
    # Group chunks by source
    sources = {}
    for chunk in chunks:
        source = chunk.metadata.get('source', 'Unknown')
        if source not in sources:
            sources[source] = 0
        sources[source] += 1
    
    logger.info("\n" + "="*60)
    logger.info("INDEX SUMMARY")
    logger.info("="*60)
    logger.info(f"Total chunks: {len(chunks)}")
    logger.info(f"Total documents: {len(sources)}")
    logger.info("\nChunks per document:")
    for source, count in sorted(sources.items()):
        logger.info(f"  - {source}: {count} chunks")
    logger.info("="*60 + "\n")


def main():
    """Main function to create document index"""
    try:
        logger.info("Starting document indexing process...")
        logger.info(f"Looking for documents in: {os.path.abspath(DATA_DIR)}")
        
        # Step 1: Load documents
        documents = load_documents()
        
        if not documents:
            logger.error("No documents found or loaded. Exiting.")
            sys.exit(1)
        
        logger.info(f"✓ Loaded {len(documents)} documents")
        
        # Step 2: Create chunks with metadata
        chunks = create_document_chunks(documents)
        
        if not chunks:
            logger.error("Failed to create document chunks. Exiting.")
            sys.exit(1)
        
        logger.info(f"✓ Created {len(chunks)} chunks from {len(documents)} documents")
        
        # Step 3: Create FAISS index
        db = create_index(chunks)
        
        if not db:
            logger.error("Failed to create index. Exiting.")
            sys.exit(1)
        
        logger.info("✓ FAISS index created successfully")
        
        # Step 4: Save index
        success = save_index(db)
        
        if not success:
            logger.error("Failed to save index. Exiting.")
            sys.exit(1)
        
        # Step 5: Print summary
        print_summary(chunks)
        
        logger.info("✅ Indexing process completed successfully!")
        
    except KeyboardInterrupt:
        logger.warning("\n⚠ Indexing interrupted by user")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Unexpected error: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()